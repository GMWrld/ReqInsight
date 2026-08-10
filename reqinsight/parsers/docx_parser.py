from pathlib import Path

from docx import Document as DocxDocument

from reqinsight.models.document import Document
from reqinsight.parsers.base_parser import BaseDocumentParser


class DOCXParser(BaseDocumentParser):
    """Parser for Microsoft Word DOCX files."""

    def parse(self, file_path: str) -> Document:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document not found: {file_path}"
            )

        docx_document = DocxDocument(file_path)

        paragraphs = []

        for paragraph in docx_document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        text = "\n".join(paragraphs)

        document = Document(file_path)
        document.text = text

        return document